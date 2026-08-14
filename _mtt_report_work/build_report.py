import json, shutil, os, stat
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r"D:\01-科院-工作\01-应龙\005-原型\应龙-0727资源权限")
WORK=ROOT/'_mtt_report_work'
REF=Path(r"D:\01-科院-工作\01-应龙\012-大模型测评\BW1000测试报告-20251201.docx")
OUT=ROOT/'MTT-S5000测试报告-20260729.docx'
data=json.loads((WORK/'analysis.json').read_text(encoding='utf-8'))

if OUT.exists(): os.chmod(OUT, os.stat(OUT).st_mode | stat.S_IWRITE)
shutil.copy2(REF,OUT)
os.chmod(OUT, os.stat(OUT).st_mode | stat.S_IWRITE)
doc=Document(OUT)
body=doc._element.body
sectPr=body.sectPr
for child in list(body):
    if child is not sectPr: body.remove(child)

def set_run_font(run, name='宋体', size=10.5, bold=None, color=None):
    run.font.name=name; run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)
    rPr=run._element.get_or_add_rPr(); rf=rPr.rFonts
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.insert(0,rf)
    for k in ('ascii','hAnsi','eastAsia','cs'): rf.set(qn('w:'+k),name)

for style_name,size,bold in [('Normal',10.5,False),('Heading 1',16,True),('Heading 2',14,True),('Heading 3',12,True)]:
    s=doc.styles[style_name]; s.font.name='宋体'; s.font.size=Pt(size); s.font.bold=bold
    s._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    pf=s.paragraph_format; pf.space_after=Pt(6); pf.line_spacing=1.25
for hn in ('Heading 1','Heading 2','Heading 3'):
    doc.styles[hn].font.color.rgb=RGBColor(0,0,0)

def add_text(text='',style='Normal',align=None,bold=False,size=None,indent=True):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.25
    if indent and style=='Normal': p.paragraph_format.first_line_indent=Cm(0.74)
    if align is not None: p.alignment=align
    role_size={'Normal':10.5,'Heading 1':16,'Heading 2':14,'Heading 3':12}.get(style,10.5)
    r=p.add_run(text); set_run_font(r,size=size or role_size,bold=(bold or style.startswith('Heading')))
    return p

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def margins(cell,top=80,start=100,bottom=80,end=100):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=m.find(qn('w:'+tag))
        if el is None: el=OxmlElement('w:'+tag); m.append(el)
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')

def add_table(headers,rows,widths=None,font=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=str(h); shade(c,'D9E2F3'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(c)
    for row in rows:
        cells=t.add_row().cells
        for j,v in enumerate(row):
            cells[j].text='' if v is None else str(v); cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cells[j])
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Cm(w)
    for ri,row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: set_run_font(r,size=font,bold=(ri==0))
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def caption(text): return add_text(text,align=WD_ALIGN_PARAGRAPH.CENTER,indent=False,size=10)
def fmt(x,d=2): return '' if x is None else f'{x:.{d}f}'

# Cover
for _ in range(6): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(18)
r=p.add_run('MTT-S5000 测试报告'); set_run_font(r,'黑体',26,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('大模型与常用小模型推理性能测试'); set_run_font(r,'宋体',14,False,(89,89,89))
for _ in range(9): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('测试数据汇总日期：2026年7月29日'); set_run_font(r,'宋体',11)
doc.add_page_break()

add_text('1 概述','Heading 1',indent=False)
add_text('本次测试主要针对 MTT-S5000 平台开展推理性能验证，覆盖 Qwen3-32B、Qwen3.5-27B 两款大模型的在线与离线推理场景，以及 Qwen3-Embedding-0.6B、Qwen3-Reranker-0.6B 两款常用小模型。报告依据 6 份原始 CSV 数据汇总，重点分析不同输入/输出长度和并发配置下的首 token 时延、非首 token 时延、吞吐量与请求时延。')
add_text('在线场景沿用参考报告口径：输入长度不超过 4096 时，以最大首 token 时延不超过 2 s 为目标；输入长度达到 8192 及以上时，以不超过 20 s 为目标，并在满足门槛的结果中选取最大并发数。离线场景不强调交互时延，按同一输入/输出组合下总吞吐量最大的配置汇总。')

add_text('2 测试环境','Heading 1',indent=False)
add_text('2.1 硬件环境','Heading 2',indent=False)
caption('表1 硬件环境')
add_table(['序号','设备','说明'],[['1','MTT-S5000 服务器','测试数据标识为 Hygon 2_1；CPU、内存、加速卡数量及互联配置未随 CSV 提供，需依据测试现场记录补充。']], [1.2,3.4,12.0],9)
add_text('2.2 软件环境','Heading 2',indent=False)
caption('表2 软件与模型环境')
add_table(['序号','项目','说明'],[
 ['1','推理模式','在线推理、离线批处理'],['2','大模型','Qwen3-32B、Qwen3.5-27B'],['3','小模型','Qwen3-Embedding-0.6B、Qwen3-Reranker-0.6B'],['4','软件栈','操作系统、驱动、容器、推理框架及精度配置未随 CSV 提供，需补充。']], [1.2,4.2,11.2],9)
add_text('2.3 环境差异说明','Heading 2',indent=False)
add_text('本报告仅对所提供测试数据进行汇总，不对未提供的硬件规格、软件版本、量化精度、批处理策略及调优参数作推断。后续若与其他平台对比，应确保模型权重、精度、输入数据、并发策略和统计口径一致。')

add_text('3 测试结果','Heading 1',indent=False)
add_text('3.1 大模型推理','Heading 2',indent=False)
add_text('3.1.1 在线推理','Heading 3',indent=False)
add_text('在线推理主要关注交互体验。下表优先列出满足首 token 时延门槛的最大并发配置；若某场景无任何配置满足门槛，则列出该场景首 token 时延最小的实测配置，并标记“未达标”。')
table_no=3
for model_key,title in [('qwen3-32b','Qwen3-32B'),('qwen3.5-27b','Qwen3.5-27B')]:
    rows=[]
    for x in data['models'][model_key]['online']:
        rows.append([x['input'],x['output'],x['concurrency'],fmt(x['ttft_max_s']),fmt(x['itl_ms']),fmt(x['throughput_avg']),fmt(x['per_req_avg']),'达标' if x['qualified'] else '未达标'])
    caption(f'表{table_no} MTT-S5000 {title} 在线推理性能汇总'); table_no+=1
    add_table(['输入','输出','并发','最大首tk(s)','平均非首tk(ms)','总吞吐(tk/s)','单请求吞吐','门槛'],rows,[1.3,1.3,1.1,2.0,2.3,2.1,2.0,1.4],7.7)

q32=data['models']['qwen3-32b']['online']; q35=data['models']['qwen3.5-27b']['online']
add_text('结果显示，两款模型在 512、2048、8192 输入及 Qwen3.5-27B 的 20000 输入场景下均可找到满足既定首 token 门槛的并发配置。4096 输入/512 输出场景未出现最大首 token 时延不超过 2 s 的配置：Qwen3-32B 的最小实测最大首 token 时延为 %.2f s，Qwen3.5-27B 为 %.2f s，应作为在线体验优化重点。' % (next(x for x in q32 if x['input']==4096)['ttft_max_s'],next(x for x in q35 if x['input']==4096)['ttft_max_s']))

add_text('3.1.2 离线批处理','Heading 3',indent=False)
add_text('离线场景以同一输入/输出组合下总吞吐量最大的配置作为推荐结果。该口径适用于批量数据分析、离线生成等不以单请求响应时间为首要目标的任务。')
for model_key,title in [('qwen3-32b','Qwen3-32B'),('qwen3.5-27b','Qwen3.5-27B')]:
    rows=[]
    for x in data['models'][model_key]['offline']:
        rows.append([x['input'],x['output'],x['concurrency'],fmt(x['latency_avg_s'],4),fmt(x['throughput_avg'],4),fmt(x['per_req_avg'],4),fmt(x['generation_speed'],4)])
    caption(f'表{table_no} MTT-S5000 {title} 离线推理性能汇总'); table_no+=1
    add_table(['输入','输出','并发','并发时延(s)','总吞吐(tk/s)','单并发吞吐','生成速度'],rows,[1.5,1.5,1.2,2.5,2.5,2.5,2.2],8)

for model_key,title in [('qwen3-32b','Qwen3-32B'),('qwen3.5-27b','Qwen3.5-27B')]:
    peak=max(data['models'][model_key]['offline'],key=lambda x:x['throughput_avg'])
    add_text(f'{title} 离线测试的峰值总吞吐为 {peak["throughput_avg"]:.2f} tokens/s，对应输入 {peak["input"]}、输出 {peak["output"]}、并发 {peak["concurrency"]}。')

add_text('3.2 常用小模型推理','Heading 2',indent=False)
e=data['small_models']['embedding'][0]; rr=data['small_models']['reranker'][0]
caption(f'表{table_no} Qwen3-Embedding-0.6B 推理性能汇总'); table_no+=1
add_table(['输入字符','并发','请求数','成功/失败','吞吐(req/s)','平均时延(s)','P50/P90/P99(s)','向量维度'],[[e['input_chars'],e['concurrency'],e['requests'],f"{e['success']}/{e['failed']}",fmt(float(e['throughput_req_s']),3),fmt(float(e['avg_latency_s']),3),f"{float(e['p50_latency_s']):.3f}/{float(e['p90_latency_s']):.3f}/{float(e['p99_latency_s']):.3f}",e['embedding_dim']]],[1.6,1.1,1.3,1.6,2.0,1.8,3.2,1.7],8)
caption(f'表{table_no} Qwen3-Reranker-0.6B 推理性能汇总'); table_no+=1
add_table(['Query/文档tk','候选文档','并发','请求数','成功/失败','吞吐(req/s)','候选吞吐(doc/s)','平均时延(s)','P50/P90/P99(s)'],[[f"{rr['query_tokens_target']}/{rr['document_tokens_target']}",rr['documents_per_request'],rr['concurrency'],rr['requests'],f"{rr['success']}/{rr['failed']}",fmt(float(rr['throughput_req_s']),3),fmt(float(rr['throughput_candidate_doc_s']),3),fmt(float(rr['avg_latency_s']),3),f"{float(rr['p50_latency_s']):.2f}/{float(rr['p90_latency_s']):.2f}/{float(rr['p99_latency_s']):.2f}"]],[1.7,1.3,1.0,1.1,1.4,1.6,2.1,1.7,2.8],7.7)
add_text('Embedding 测试 100 个请求全部成功，吞吐为 83.29 req/s，平均时延 0.186 s。Reranker 测试同样 100% 成功，在每请求 30 个候选文档、并发 16 的条件下，吞吐为 0.738 req/s（22.15 candidate-doc/s），平均时延 21.17 s。')

add_text('4 总结','Heading 1',indent=False)
add_text('MTT-S5000 已完成 Qwen3-32B、Qwen3.5-27B 的在线与离线推理测试，并完成 Embedding、Reranker 小模型验证；所提供结果中未出现失败请求记录。大模型在线场景在短输入、常见 2048 输入和 8192 长输入场景均能形成满足既定首 token 门槛的配置，但 4096 输入/512 输出场景未达到 2 s 门槛，需要从调度、批处理、KV Cache、算子与通信等环节进一步分析。')
add_text('离线场景应以吞吐优先配置部署，并根据实际业务输入分布选择并发。当前数据能够说明模型在 MTT-S5000 平台上的可运行性和基础性能，但由于缺少完整软硬件环境、精度配置及对照平台数据，本报告不作跨硬件优劣结论。建议补齐环境信息后，增加稳定性、显存占用、功耗、整机多实例及对照平台测试。')

# Footer page number
for section in doc.sections:
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.clear(); r=p.add_run('第 '); set_run_font(r,size=9)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)
    r=p.add_run(' 页'); set_run_font(r,size=9)
    section.different_first_page_header_footer=True

doc.core_properties.title='MTT-S5000 测试报告'
doc.core_properties.subject='大模型与常用小模型推理性能测试汇总'
doc.core_properties.author=''
doc.save(OUT)
print('report_saved')
