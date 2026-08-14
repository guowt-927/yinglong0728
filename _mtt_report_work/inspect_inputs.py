import csv, json, os, sys
from pathlib import Path
from docx import Document

ref = Path(r"D:\01-科院-工作\01-应龙\012-大模型测评\BW1000测试报告-20251201.docx")
csv_dir = Path(r"C:\Users\wanting\Desktop\新建文件")
out = Path(__file__).parent

doc = Document(ref)
data = {
    "sections": [{
        "page_width": s.page_width, "page_height": s.page_height,
        "top_margin": s.top_margin, "bottom_margin": s.bottom_margin,
        "left_margin": s.left_margin, "right_margin": s.right_margin,
        "header_distance": s.header_distance, "footer_distance": s.footer_distance,
    } for s in doc.sections],
    "paragraphs": [], "tables": []
}
for i,p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        data["paragraphs"].append({"i":i,"style":p.style.name if p.style else None,"text":text})
for ti,t in enumerate(doc.tables):
    rows=[]
    for r in t.rows:
        rows.append([c.text.strip().replace("\n"," | ") for c in r.cells])
    data["tables"].append({"i":ti,"rows":rows})

csvs={}
for path in sorted(csv_dir.glob("*.csv")):
    raw=path.read_bytes()
    enc=None
    text=None
    for e in ("utf-8-sig","utf-8","gb18030","utf-16"):
        try:
            text=raw.decode(e); enc=e; break
        except UnicodeDecodeError: pass
    if text is None: continue
    sample=text[:4096]
    try: dialect=csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error: dialect=csv.excel
    rows=list(csv.reader(text.splitlines(), dialect))
    csvs[path.name]={"encoding":enc,"rows":rows}

(out/"inputs.json").write_text(json.dumps({"doc":data,"csvs":csvs},ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"paragraphs":len(data["paragraphs"]),"tables":len(data["tables"]),"csvs":{k:len(v["rows"]) for k,v in csvs.items()}},ensure_ascii=False))
